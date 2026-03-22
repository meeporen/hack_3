import base64
import json
import os
import sys
import tempfile

import pandas as pd
from docx import Document
from openai import OpenAI

_MODEL_VISION = "qwen/qwen2.5-vl-72b-instruct"


def generate_schema_hint(filepath: str, api_key: str | None = None) -> dict:
    if api_key is None:
        from src.config import settings
        api_key = settings.OPENROUTER_API_KEY

    print(f"\n[docx_parser] ── старт ─────────────────────────────────────")
    print(f"[docx_parser] filepath : {filepath}")

    print(f"\n[docx_parser] ── шаг 1: python-docx ───────────────────────")
    schema = _try_python_docx(filepath)
    if schema is not None:
        print(f"[docx_parser] ✓ структурированная таблица найдена, vision не нужен")
        return schema

    print(f"[docx_parser] ✗ таблица не найдена — переходим в vision pipeline")
    return _vision_pipeline(filepath, api_key)


def _try_python_docx(filepath: str) -> dict | None:
    try:
        doc = Document(filepath)
    except Exception as e:
        print(f"[docx_parser] python-docx ошибка: {e}")
        return None

    if not doc.tables:
        print(f"[docx_parser] python-docx: таблиц не найдено")
        return None

    print(f"[docx_parser] python-docx: найдено таблиц: {len(doc.tables)}")
    table = max(doc.tables, key=lambda t: len(t.rows) * len(t.columns))
    df = _df_from_docx_table(table)

    if df.empty:
        print(f"[docx_parser] python-docx: таблица пустая после парсинга")
        return None

    print(f"[docx_parser] python-docx: таблица {len(df)}x{len(df.columns)}")

    csv_bytes = df.to_csv(sep=";", index=False).encode("utf-8")

    return {
        "file_type": "csv",
        "separator": ";",
        "row_count": len(df),
        "col_count": len(df.columns),
        "columns": _build_columns(df),
        "_csv_bytes_b64": base64.b64encode(csv_bytes).decode(),
        "_original_type": "docx",
        "_strategy": "python-docx",
    }


def _vision_pipeline(filepath: str, api_key: str) -> dict:
    try:
        from parsers.pdf_parser import _rasterize_pdf, _strip_markdown, _PROMPT_HEADERS, _PROMPT_DATA, _parse_json_rows
        from parsers.csv_parser import _generate_schema_hint
    except ModuleNotFoundError:
        from pdf_parser import _rasterize_pdf, _strip_markdown, _PROMPT_HEADERS, _PROMPT_DATA, _parse_json_rows
        from csv_parser import _generate_schema_hint

    client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)

    print(f"\n[docx_parser] ── шаг 2а: конвертация DOCX → PDF ───────────")
    pdf_path = _docx_to_pdf(filepath)
    print(f"[docx_parser] PDF: {pdf_path}")

    print(f"\n[docx_parser] ── шаг 2б: рендер страниц PDF ────────────────")
    page_images = _rasterize_pdf(pdf_path, dpi=120)
    print(f"[docx_parser] страниц отрендерено: {len(page_images)}")

    try:
        os.unlink(pdf_path)
    except OSError:
        pass

    all_rows: list[dict] = []
    headers: list[str] | None = None
    total_in, total_out = 0, 0

    for i, img_path in enumerate(page_images):
        page_num = i + 1
        print(f"\n[docx_parser] ── страница {page_num}/{len(page_images)} ─────────────────────")

        b64, mime = _img_to_b64(img_path)

        if headers is None:
            print(f"[docx_parser] запрос заголовков...")
            r_h = client.chat.completions.create(
                model=_MODEL_VISION,
                temperature=0,
                messages=[{"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                    {"type": "text", "text": _PROMPT_HEADERS},
                ]}],
            )
            u = r_h.usage
            total_in += u.prompt_tokens; total_out += u.completion_tokens
            print(f"[docx_parser] токены (заголовки): in={u.prompt_tokens} out={u.completion_tokens}")
            raw_h = _strip_markdown(r_h.choices[0].message.content.strip())
            try:
                headers = json.loads(raw_h)
                print(f"[docx_parser] заголовки ({len(headers)}): {headers}")
            except json.JSONDecodeError:
                print(f"[docx_parser] WARN: не удалось распарсить заголовки: {raw_h!r}")
                continue

        numbered = "\n".join(f"{i+1}. {h}" for i, h in enumerate(headers))
        prompt_data = _PROMPT_DATA.format(numbered_headers=numbered, count=len(headers))

        r_d = client.chat.completions.create(
            model=_MODEL_VISION,
            temperature=0,
            messages=[{"role": "user", "content": [
                {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                {"type": "text", "text": prompt_data},
            ]}],
        )
        u = r_d.usage
        total_in += u.prompt_tokens; total_out += u.completion_tokens
        print(f"[docx_parser] токены (данные): in={u.prompt_tokens} out={u.completion_tokens}")
        raw_d = _strip_markdown(r_d.choices[0].message.content.strip())
        print(f"[docx_parser] raw данные ({len(raw_d)} chars): {raw_d[:300]}")

        page_rows = _parse_json_rows(raw_d)
        print(f"[docx_parser] строк на странице: {len(page_rows)}")
        all_rows.extend(page_rows)

        try:
            os.unlink(img_path)
        except OSError:
            pass

    print(f"\n[docx_parser] ── токены итого ───────────────────────────────")
    print(f"[docx_parser] input:  {total_in}")
    print(f"[docx_parser] output: {total_out}")
    print(f"[docx_parser] total:  {total_in + total_out}")
    print(f"\n[docx_parser] ── шаг 3: итого строк: {len(all_rows)} ─────────")

    df = pd.DataFrame(all_rows)
    csv_bytes = df.to_csv(sep=";", index=False).encode("utf-8")

    lines = csv_bytes.decode("utf-8").splitlines()
    print(f"\n[docx_parser] ── CSV ({len(lines)} строк) ─────────────────────")
    for line in lines[:20]:
        print(f"[docx_parser]   {line}")
    if len(lines) > 20:
        print(f"[docx_parser]   ... ({len(lines) - 20} more lines)")

    with tempfile.NamedTemporaryFile(suffix=".csv", delete=False, mode="wb") as tmp:
        tmp.write(csv_bytes)
        tmp_path = tmp.name

    try:
        schema = _generate_schema_hint(tmp_path)
    finally:
        os.unlink(tmp_path)

    schema["_csv_bytes_b64"] = base64.b64encode(csv_bytes).decode()
    schema["_original_type"] = "docx"
    schema["_strategy"] = "vision"
    schema["_vision_prompt_tokens"] = total_in
    schema["_vision_completion_tokens"] = total_out
    return schema


def _docx_to_pdf(filepath: str) -> str:
    """Конвертирует DOCX в PDF, возвращает путь к временному PDF."""
    out_path = tempfile.mktemp(suffix=".pdf")
    try:
        from docx2pdf import convert
        convert(filepath, out_path)
        return out_path
    except Exception:
        pass

    import subprocess
    out_dir = tempfile.mkdtemp()
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, filepath],
        check=True, capture_output=True,
    )
    name = os.path.splitext(os.path.basename(filepath))[0] + ".pdf"
    return os.path.join(out_dir, name)


def _img_to_b64(path: str) -> tuple[str, str]:
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode(), "image/png"


def _df_from_docx_table(table) -> pd.DataFrame:
    rows = [[cell.text.replace('\n', ' ').strip() for cell in row.cells] for row in table.rows]
    if not rows:
        return pd.DataFrame()

    headers = rows[0]
    data = rows[1:]

    seen = {}
    new_cols = []
    for col in headers:
        if col in seen:
            seen[col] += 1
            new_cols.append(f"{col}_{seen[col]}")
        else:
            seen[col] = 0
            new_cols.append(col)

    df = pd.DataFrame(data, columns=new_cols)
    df = df.replace("", pd.NA).dropna(how="all").reset_index(drop=True)
    df = df.apply(lambda col: pd.to_numeric(col, errors="coerce").fillna(col))
    return df


def _build_columns(df: pd.DataFrame) -> list:
    columns = []
    for col in df.columns:
        series = df[col]
        non_null = series.dropna()
        sample = non_null.unique()[0] if len(non_null) > 0 else None
        if hasattr(sample, "item"):
            sample = sample.item()
        columns.append({
            "name": str(col),
            "dtype": _map_dtype(series),
            "sample": str(sample) if sample is not None else None,
            "has_nulls": bool(series.isna().any()),
        })
    return columns


def _map_dtype(series: pd.Series) -> str:
    if pd.api.types.is_bool_dtype(series):            return "bool"
    if pd.api.types.is_integer_dtype(series):         return "int64"
    if pd.api.types.is_float_dtype(series):           return "float64"
    if pd.api.types.is_datetime64_any_dtype(series):  return "datetime"
    if pd.api.types.is_timedelta64_dtype(series):     return "timedelta"
    if isinstance(series.dtype, pd.CategoricalDtype): return "category"
    return "str"


if __name__ == "__main__":
    path = sys.argv[1] if len(sys.argv) > 1 else input("Путь к DOCX: ")
    try:
        result = generate_schema_hint(path)
        result.pop("_csv_bytes_b64", None)
        print(json.dumps(result, ensure_ascii=False, indent=4))
    except Exception:
        import traceback
        traceback.print_exc()
